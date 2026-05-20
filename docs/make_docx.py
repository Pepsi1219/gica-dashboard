"""
Generate CSA Manager user guide as a DOCX.
Run: python make_docx.py
Output: csa-manager-guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ─────────────────────────────────────────────────────────────────
DEEP   = RGBColor(0x0B, 0x4E, 0x63)
BLUE   = RGBColor(0x3B, 0x8D, 0xC3)
INK    = RGBColor(0x0E, 0x28, 0x33)
MUTED  = RGBColor(0x5A, 0x70, 0x80)
OK     = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xEF, 0x44, 0x44)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Tahoma"

doc = Document()

# Set default font + size for normal style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
# Asian font fallback
rpr = style.element.get_or_add_rPr()
fnt = rpr.find(qn("w:rFonts"))
if fnt is None:
    fnt = OxmlElement("w:rFonts")
    rpr.append(fnt)
fnt.set(qn("w:eastAsia"), FONT)
fnt.set(qn("w:ascii"), FONT)
fnt.set(qn("w:hAnsi"), FONT)
fnt.set(qn("w:cs"), FONT)

# narrow margins for compact look
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)


def shade_cell(cell, hex6):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tc_pr.append(shd)


def add_heading(text, *, size=20, color=DEEP, bold=True, space_before=14, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    # also set east asia font for thai
    rpr = r._r.get_or_add_rPr()
    fnt_el = OxmlElement("w:rFonts")
    fnt_el.set(qn("w:eastAsia"), FONT)
    fnt_el.set(qn("w:ascii"), FONT)
    fnt_el.set(qn("w:hAnsi"), FONT)
    fnt_el.set(qn("w:cs"), FONT)
    rpr.append(fnt_el)
    return p


def add_para(text, *, size=11, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    rpr = r._r.get_or_add_rPr()
    fnt_el = OxmlElement("w:rFonts")
    fnt_el.set(qn("w:eastAsia"), FONT)
    fnt_el.set(qn("w:ascii"), FONT)
    fnt_el.set(qn("w:hAnsi"), FONT)
    fnt_el.set(qn("w:cs"), FONT)
    rpr.append(fnt_el)
    return p


def add_bullet(text, *, size=11, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    rpr = r._r.get_or_add_rPr()
    fnt_el = OxmlElement("w:rFonts")
    fnt_el.set(qn("w:eastAsia"), FONT)
    fnt_el.set(qn("w:ascii"), FONT)
    fnt_el.set(qn("w:hAnsi"), FONT)
    fnt_el.set(qn("w:cs"), FONT)
    rpr.append(fnt_el)
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════
title_table = doc.add_table(rows=1, cols=1)
title_table.autofit = False
title_table.columns[0].width = Cm(17)
title_cell = title_table.rows[0].cells[0]
shade_cell(title_cell, "0B4E63")
# Padding
tc_pr = title_cell._tc.get_or_add_tcPr()
mar = OxmlElement("w:tcMar")
for side in ("top", "left", "bottom", "right"):
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:w"), "300")
    el.set(qn("w:type"), "dxa")
    mar.append(el)
tc_pr.append(mar)

# title text
p1 = title_cell.paragraphs[0]
r1 = p1.add_run("คู่มือการใช้งาน CSA Manager")
r1.font.name = FONT; r1.font.size = Pt(24); r1.font.bold = True
r1.font.color.rgb = WHITE
rpr = r1._r.get_or_add_rPr()
fnt_el = OxmlElement("w:rFonts")
for a in ("eastAsia", "ascii", "hAnsi", "cs"):
    fnt_el.set(qn(f"w:{a}"), FONT)
rpr.append(fnt_el)

p2 = title_cell.add_paragraph()
r2 = p2.add_run("โหมดผู้บริหาร · ดูข้อมูลอย่างเดียว · ไม่ต้องใช้รหัส")
r2.font.name = FONT; r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0xCA, 0xDC, 0xFC)
rpr2 = r2._r.get_or_add_rPr()
fnt_el2 = OxmlElement("w:rFonts")
for a in ("eastAsia", "ascii", "hAnsi", "cs"):
    fnt_el2.set(qn(f"w:{a}"), FONT)
rpr2.append(fnt_el2)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — LOGIN
# ══════════════════════════════════════════════════════════════════════════
add_heading("1. วิธีเข้าสู่ระบบ", size=16)
add_para("เข้าสู่ระบบได้ใน 3 ขั้นตอนสั้นๆ — ไม่ต้องกรอก username หรือ password",
         color=MUTED, italic=True, space_after=8)

steps_tbl = doc.add_table(rows=3, cols=2)
steps_tbl.autofit = False
steps_tbl.columns[0].width = Cm(1.2)
steps_tbl.columns[1].width = Cm(15.5)

steps_data = [
    ("1", "เปิดเว็บไซต์ของระบบ",
     "ใส่ที่อยู่เว็บที่ทีมไอทีส่งให้ในเบราว์เซอร์ (Chrome / Edge แนะนำ)"),
    ("2", "คลิกปุ่ม “สำหรับ CSA Manager”",
     "ที่หน้าเข้าสู่ระบบ จะเห็นปุ่มนี้อยู่ใต้ปุ่ม Sign in with Microsoft"),
    ("3", "เข้าระบบสำเร็จ",
     "ระบบจะพาเข้าสู่หน้าแรกทันที จะเห็นแถบสีฟ้าด้านขวาบนเขียนว่า “CSA Manager (อ่านอย่างเดียว)”"),
]

for i, (num, ttl, desc) in enumerate(steps_data):
    row = steps_tbl.rows[i]
    # number cell
    c0 = row.cells[0]
    shade_cell(c0, "0B4E63")
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(num)
    r.font.name = FONT; r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = WHITE
    rpr = r._r.get_or_add_rPr()
    fnt_el = OxmlElement("w:rFonts")
    for a in ("eastAsia", "ascii", "hAnsi", "cs"):
        fnt_el.set(qn(f"w:{a}"), FONT)
    rpr.append(fnt_el)
    # content cell
    c1 = row.cells[1]
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pt = c1.paragraphs[0]
    rt = pt.add_run(ttl)
    rt.font.name = FONT; rt.font.size = Pt(13); rt.font.bold = True
    rt.font.color.rgb = INK
    rpr3 = rt._r.get_or_add_rPr()
    fnt_el3 = OxmlElement("w:rFonts")
    for a in ("eastAsia", "ascii", "hAnsi", "cs"):
        fnt_el3.set(qn(f"w:{a}"), FONT)
    rpr3.append(fnt_el3)

    pd = c1.add_paragraph()
    pd.paragraph_format.space_after = Pt(0)
    rd = pd.add_run(desc)
    rd.font.name = FONT; rd.font.size = Pt(10)
    rd.font.color.rgb = MUTED
    rpr4 = rd._r.get_or_add_rPr()
    fnt_el4 = OxmlElement("w:rFonts")
    for a in ("eastAsia", "ascii", "hAnsi", "cs"):
        fnt_el4.set(qn(f"w:{a}"), FONT)
    rpr4.append(fnt_el4)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — HOME SCREEN
# ══════════════════════════════════════════════════════════════════════════
add_heading("2. หน้าจอแรก", size=16, space_before=18)

add_heading("2.1 เลือกหน่วยงาน (ส่วนบน)", size=12, color=BLUE, space_before=6, space_after=2)
add_para("เห็นกล่องหน่วยงานทั้งหมด: G1, G2, G3, G4, TRM, EA — คลิกกล่องไหน "
         "ระบบจะพาไปหน้า Dashboard ของหน่วยงานนั้น")

add_heading("2.2 ภาพรวมทั้งหมด (ส่วนล่าง)", size=12, color=BLUE, space_before=8, space_after=2)
add_para("สรุปข้อมูลของทุกหน่วยงานในที่เดียว ประกอบด้วย:")
add_bullet("ตารางสรุป: ทั้งหมด / ฝึกเสร็จ / กำลังฝึก / ลาออก / โอนย้าย ของแต่ละหน่วยงาน")
add_bullet("Data Analytics: กราฟวงกลม 4 ตัว + การ์ดสถิติ + กราฟแนวโน้มรายเดือน")
add_bullet("ฟิลเตอร์ปี: ดูเฉพาะปี หรือ “ทุกปี” เพื่อดูภาพรวม — เปลี่ยนแล้วทุกอย่างอัพเดททันที")
add_bullet("ปุ่ม PNG / PDF: บันทึกหน้าเป็นรูปภาพหรือไฟล์ PDF ไว้ดูทีหลัง "
           "(ไฟล์จะมีปีที่เลือกแสดงด้วย)")


# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — DEPARTMENT DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
add_heading("3. หน้า Dashboard ของหน่วยงาน", size=16, space_before=18)
add_para("เข้าได้โดยคลิกกล่องหน่วยงานที่หน้าแรก จะพบส่วนต่างๆ ดังนี้",
         color=MUTED, italic=True)

add_bullet("การ์ดสรุปด้านบน 5 ใบ — ดูจำนวนพนักงานในแต่ละสถานะของหน่วยงานนี้: "
           "ทั้งหมด, ฝึกเสร็จ (ตรงเวลา/เกินกำหนด), กำลังฝึก (พื้นฐาน/ขั้นตอน), "
           "การลาออก (3 ระยะ), การโอนย้าย (3 ระยะ)")
add_bullet("ค้นหารหัสพนักงาน — พิมพ์รหัสในช่องค้นหา จะเปิดหน้าต่างแสดงรายละเอียดทั้งหมดของพนักงานคนนั้น")
add_bullet("ฟิลเตอร์ ปี / สถานะ — เลือกได้ทั้งสองอัน เพื่อให้ตารางเหลือเฉพาะรายการที่ตรงตามที่เลือก")
add_bullet("ตารางพนักงาน — เลื่อนซ้าย-ขวาเพื่อดูครบทุกคอลัมน์ (ชื่อ, เกรด, วันที่ฝึก, สถานะจริง, ฯลฯ)")


# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════
add_heading("4. สิทธิ์การใช้งาน", size=16, space_before=18)

perm_tbl = doc.add_table(rows=1, cols=2)
perm_tbl.autofit = False
perm_tbl.columns[0].width = Cm(8.5)
perm_tbl.columns[1].width = Cm(8.5)

# headers
hdr_can = perm_tbl.rows[0].cells[0]
hdr_cant = perm_tbl.rows[0].cells[1]
shade_cell(hdr_can, "10B981")
shade_cell(hdr_cant, "EF4444")

for cell, txt in ((hdr_can, "✓  สิ่งที่ทำได้"), (hdr_cant, "✗  สิ่งที่ทำไม่ได้")):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.name = FONT; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = WHITE
    rpr = r._r.get_or_add_rPr()
    fnt_el = OxmlElement("w:rFonts")
    for a in ("eastAsia", "ascii", "hAnsi", "cs"):
        fnt_el.set(qn(f"w:{a}"), FONT)
    rpr.append(fnt_el)

# Data rows
can_items = [
    "ดูตาราง “ภาพรวมทั้งหมด” ของทุกหน่วยงาน",
    "ดูสถิติและกราฟใน Data Analytics",
    "เปลี่ยนฟิลเตอร์ปีและสถานะ",
    "เปิด Dashboard ของแต่ละหน่วยงาน",
    "ค้นหารหัสพนักงาน → ดูรายละเอียด",
    "บันทึกหน้าเป็น PNG หรือ PDF",
    "เปลี่ยนภาษา (TH / EN / LO / VI)",
    "เปลี่ยนธีม สว่าง / มืด",
]
cant_items = [
    "ลงทะเบียนพนักงานใหม่ (ปุ่ม “ลงทะเบียน” ถูกซ่อน)",
    "แก้ไขข้อมูลในตาราง (ปุ่ม ✏ ถูกซ่อน)",
    "บันทึก/อัพเดทข้อมูลในหน้าต่างค้นหา",
    "เปิดหน้า Admin (เฟือง ⚙ ถูกซ่อน)",
    "เพิ่ม/ลบวันหยุดพิเศษ",
]

# add cells row by row by adding new rows
max_rows = max(len(can_items), len(cant_items))
for i in range(max_rows):
    row = perm_tbl.add_row()
    for col_idx, items in enumerate((can_items, cant_items)):
        c = row.cells[col_idx]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if i < len(items):
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            symbol = "✓ " if col_idx == 0 else "✗ "
            sym_color = OK if col_idx == 0 else DANGER
            rs = p.add_run(symbol)
            rs.font.name = FONT; rs.font.size = Pt(11); rs.font.bold = True
            rs.font.color.rgb = sym_color
            rpr_s = rs._r.get_or_add_rPr()
            fnt_el_s = OxmlElement("w:rFonts")
            for a in ("eastAsia", "ascii", "hAnsi", "cs"):
                fnt_el_s.set(qn(f"w:{a}"), FONT)
            rpr_s.append(fnt_el_s)

            r = p.add_run(items[i])
            r.font.name = FONT; r.font.size = Pt(11)
            r.font.color.rgb = INK
            rpr = r._r.get_or_add_rPr()
            fnt_el = OxmlElement("w:rFonts")
            for a in ("eastAsia", "ascii", "hAnsi", "cs"):
                fnt_el.set(qn(f"w:{a}"), FONT)
            rpr.append(fnt_el)

add_para("ระบบจะซ่อนปุ่มเหล่านี้ให้อัตโนมัติเมื่อเข้าผ่านโหมด CSA Manager — "
         "ไม่ต้องกังวลว่าจะเผลอแก้ไขข้อมูล",
         color=MUTED, italic=True, space_after=4)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — TIPS
# ══════════════════════════════════════════════════════════════════════════
add_heading("5. เคล็ดลับการใช้งาน", size=16, space_before=18)

add_bullet("เปลี่ยนภาษา — คลิกธงด้านขวาบน สลับระหว่าง 🇹🇭 ไทย / 🇺🇸 English / 🇱🇦 ลาว / 🇻🇳 Tiếng Việt")
add_bullet("ธีมมืด / สว่าง — คลิกไอคอน 🌙 / ☀️ บนแถบด้านขวาบน เพื่อเปลี่ยนโหมดการแสดงผล")
add_bullet("บันทึก PNG / PDF — ในหน้า “ภาพรวมทั้งหมด” กดปุ่ม PNG หรือ PDF ที่มุมขวาบนของการ์ด "
           "ไฟล์ที่ได้จะมีปีที่เลือกไว้ติดมาด้วย — รู้ได้ว่าเป็นข้อมูลช่วงเวลาไหน")
add_bullet("ดูข้อมูลย้อนหลัง — เลือกปีในฟิลเตอร์ปี เพื่อดูข้อมูลของช่วงเวลานั้น")
add_bullet("ออกจากระบบ — คลิก “Logout” ข้างขวาของข้อความ “Signed in” บนแถบด้านบน")

# Footer note
doc.add_paragraph()  # spacer
last = doc.add_paragraph()
last.alignment = WD_ALIGN_PARAGRAPH.CENTER
last.paragraph_format.space_before = Pt(20)
rf = last.add_run("─── สอบถามเพิ่มเติม · ติดต่อทีมไอที ───")
rf.font.name = FONT; rf.font.size = Pt(10); rf.font.italic = True
rf.font.color.rgb = MUTED
rpr_f = rf._r.get_or_add_rPr()
fnt_el_f = OxmlElement("w:rFonts")
for a in ("eastAsia", "ascii", "hAnsi", "cs"):
    fnt_el_f.set(qn(f"w:{a}"), FONT)
rpr_f.append(fnt_el_f)


# ── Save ──
output = "csa-manager-guide.docx"
doc.save(output)
print(f"Saved: {output}")
